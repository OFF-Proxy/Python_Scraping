from docx import Document
from docx.shared import Inches

# 既存のWordドキュメントを読み込み
document = Document('sample.docx')

# 画像を追加するパス
image_path = './eiko.jpg'

# 画像をドキュメントに挿入
document.add_picture(image_path, width=Inches(3), height=Inches(4))

# ドキュメント内のテキストを取得し、文字数をカウント
text = ''.join([paragraph.text for paragraph in document.paragraphs])
character_count = len(text)

# 文字数を出力
print("ドキュメント内の文字数:", character_count)

# "sample_answer.docx"としてドキュメントを保存
document.save('sample_answer.docx')